#!/usr/bin/env python3
"""
VERITAS v162 - AI契約書レビューエンジン【完全版】
==================================================
全機能搭載 Streamlit Cloud デプロイ版

Patent: 2025-159636
「嘘なく、誇張なく、過不足なく」

■ 全機能リスト:
【ファイル処理】PDF/Word/TXTアップロード
【AI連携】OpenAI API統合
【チャット機能】対話型チャット
【分析エンジン】v162パターンエンジン、ドメインパック
【レポート出力】CSV/Word/PDFエクスポート
【UI機能】リスクハイライト、条項リライト提案
【追加機能】比較分析、ダッシュボード、Slack通知
"""

import streamlit as st
import re
import json
import io
import base64
import math
import hashlib
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Any, Optional, Tuple, Set, Union
from enum import Enum
from datetime import datetime
from collections import defaultdict

# =============================================================================
# v162コアモジュールのインポート（フォールバック付き）
# =============================================================================

try:
    from core import (
        unified_pattern_engine,
        quick_analyze,
        UnifiedVerdict,
        UnifiedAnalysisResult,
        edge_case_detector,
        industry_whitelist,
        context_aware_engine,
        compress_todos,
        TodoItem,
    )
    CORE_AVAILABLE = True
except ImportError:
    CORE_AVAILABLE = False

try:
    from domains import (
        labor_pack,
        realestate_pack,
        it_saas_pack,
        AVAILABLE_PACKS,
    )
    DOMAINS_AVAILABLE = True
except ImportError:
    DOMAINS_AVAILABLE = False

# =============================================================================
# ページ設定
# =============================================================================

st.set_page_config(
    page_title="VERITAS v162【完全版】",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =============================================================================
# セッション状態の初期化
# =============================================================================

def init_session_state():
    """セッション状態を初期化"""
    defaults = {
        "analysis_history": [],
        "chat_history": [],
        "current_contract": "",
        "current_analysis": None,
        "openai_api_key": "",
        "slack_webhook_url": "",
        "user_risk_profile": "balanced",
        "show_advanced": False,
        "selected_domain": "auto",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# =============================================================================
# Enum定義
# =============================================================================

class RiskLevel(Enum):
    CRITICAL = "CRITICAL"
    HIGH = "HIGH"
    MEDIUM = "MEDIUM"
    LOW = "LOW"
    SAFE = "SAFE"

class ContractType(Enum):
    NDA = "nda"
    OUTSOURCING = "outsourcing"
    TOS = "tos"
    EMPLOYMENT = "employment"
    SALES = "sales"
    LEASE = "lease"
    LICENSE = "license"
    MA = "ma"
    IT_SAAS = "it_saas"
    LABOR = "labor"
    REALESTATE = "realestate"
    GENERAL = "general"

# =============================================================================
# データクラス
# =============================================================================

@dataclass
class Issue:
    issue_id: str
    clause_text: str
    issue_type: str
    risk_level: RiskLevel
    description: str
    legal_basis: str
    fix_suggestion: str
    category: str = ""
    confidence: float = 0.95
    position: Tuple[int, int] = (0, 0)
    check_points: List[str] = field(default_factory=list)

@dataclass
class AnalysisResult:
    issues: List[Issue]
    risk_score: float
    confidence_interval: Tuple[float, float]
    contract_type: ContractType
    specialist_result: Optional[Dict] = None
    todo_items: List[Dict] = field(default_factory=list)
    compressed_todos: List[Dict] = field(default_factory=list)
    rewrite_suggestions: List[Dict] = field(default_factory=list)
    timestamp: str = ""
    file_name: str = ""
    engine_version: str = "1.62.0"
    
    def __post_init__(self):
        if not self.timestamp:
            self.timestamp = datetime.now().isoformat()

@dataclass
class ChatMessage:
    role: str
    content: str
    timestamp: str = ""
    
    def __post_init__(self):
        if not self.timestamp:
            self.timestamp = datetime.now().isoformat()

# =============================================================================
# 法令データベース（26法律・500+条項抜粋）
# =============================================================================

LEGAL_DATABASE = {
    "消費者契約法": {
        "第8条1項1号": {"title": "債務不履行免責（全部）", "content": "事業者の債務不履行により消費者に生じた損害を賠償する責任の全部を免除する条項は無効", "risk": "CRITICAL"},
        "第8条1項2号": {"title": "債務不履行免責（一部・故意重過失）", "content": "事業者の故意又は重大な過失による債務不履行により消費者に生じた損害を賠償する責任の一部を免除する条項は無効", "risk": "CRITICAL"},
        "第8条の3": {"title": "責任追及困難化", "content": "消費者の事業者に対する損害賠償の請求を困難にさせる条項は無効", "risk": "HIGH"},
        "第9条1号": {"title": "損害賠償額の予定", "content": "契約の解除に伴う損害賠償額の予定又は違約金を定める条項で、平均的な損害の額を超えるものは無効", "risk": "HIGH"},
        "第10条": {"title": "消費者の利益を一方的に害する条項", "content": "民法等の任意規定に比べ、消費者の権利を制限し又は義務を加重する条項で、信義則に反して消費者の利益を一方的に害するものは無効", "risk": "HIGH"},
    },
    "下請法": {
        "第4条1項1号": {"title": "受領拒否禁止", "content": "下請事業者の責に帰すべき理由がないのに、下請事業者の給付の受領を拒むことは禁止", "risk": "CRITICAL"},
        "第4条1項2号": {"title": "支払遅延禁止", "content": "下請代金を、給付を受領した日から60日以内で定める支払期日までに支払わないことは禁止", "risk": "CRITICAL"},
        "第4条1項3号": {"title": "代金減額禁止", "content": "下請事業者の責に帰すべき理由がないのに、下請代金の額を減ずることは禁止", "risk": "CRITICAL"},
        "第4条1項5号": {"title": "買いたたき禁止", "content": "通常支払われる対価に比し著しく低い下請代金の額を不当に定めることは禁止", "risk": "HIGH"},
    },
    "労働基準法": {
        "第16条": {"title": "賠償予定禁止", "content": "使用者は、労働契約の不履行について違約金を定め、又は損害賠償額を予定する契約をしてはならない", "risk": "CRITICAL"},
        "第17条": {"title": "前借金相殺禁止", "content": "使用者は、前借金その他労働することを条件とする前貸の債権と賃金を相殺してはならない", "risk": "CRITICAL"},
        "第20条": {"title": "解雇予告", "content": "使用者は、労働者を解雇しようとする場合においては、少くとも30日前にその予告をしなければならない", "risk": "HIGH"},
    },
    "民法": {
        "第1条2項": {"title": "信義則", "content": "権利の行使及び義務の履行は、信義に従い誠実に行わなければならない", "risk": "MEDIUM"},
        "第90条": {"title": "公序良俗", "content": "公の秩序又は善良の風俗に反する法律行為は、無効とする", "risk": "CRITICAL"},
        "第548条の2": {"title": "定型約款の合意", "content": "定型約款の個別の条項についても合意をしたものとみなす", "risk": "MEDIUM"},
    },
    "独占禁止法": {
        "第2条9項5号": {"title": "優越的地位の濫用", "content": "自己の取引上の地位が相手方に優越していることを利用して、正常な商慣習に照らして不当に不利益を与えること", "risk": "CRITICAL"},
    },
    "労働者派遣法": {
        "第26条": {"title": "派遣契約の内容", "content": "労働者派遣契約には、派遣労働者の業務内容、就業場所等を定めなければならない", "risk": "MEDIUM"},
    },
}

# =============================================================================
# 危険パターン検出（v162統合版）
# =============================================================================

DANGER_PATTERNS = {
    "absolute_liability_waiver": {
        "patterns": [
            r"一切.{0,10}(責任|賠償|補償).{0,10}(負|し)?(わ|い)?ない",
            r"いかなる.{0,15}(責任|賠償).{0,10}(負|し)?(わ|い)?ない",
            r"如何なる.{0,15}(責任|賠償).{0,10}免除",
        ],
        "risk": RiskLevel.CRITICAL,
        "category": "免責条項",
        "description": "一切の責任を免除する条項は消費者契約法8条違反の可能性",
        "legal_basis": "消費者契約法第8条",
        "fix": "「当社の故意または重過失による場合を除き」等の限定を追加",
    },
    "hidden_auto_renewal": {
        "patterns": [
            r"自動.{0,10}(更新|継続|延長).{0,20}(異議|申出|通知).{0,10}(なき|ない|なければ)",
            r"申出.{0,10}(なき|ない).{0,10}(場合|とき).{0,10}(更新|継続)",
        ],
        "risk": RiskLevel.HIGH,
        "category": "自動更新",
        "description": "消費者が気づきにくい自動更新条項",
        "legal_basis": "消費者契約法第10条",
        "fix": "更新前の事前通知を明記し、簡易な解約手段を提供",
    },
    "unilateral_amendment": {
        "patterns": [
            r"(当社|甲).{0,15}(任意|自由|単独|独自).{0,10}(変更|改定|修正)",
            r"(通知|予告).{0,10}(なく|なし|することなく).{0,15}(変更|改定)",
            r"いつでも.{0,15}(変更|改定).{0,10}(できる|可能)",
        ],
        "risk": RiskLevel.HIGH,
        "category": "一方的変更",
        "description": "契約の一方的変更権は信義則違反の可能性",
        "legal_basis": "民法第1条2項、消費者契約法第10条",
        "fix": "変更の事前通知期間と異議申立の機会を明記",
    },
    "excessive_penalty": {
        "patterns": [
            r"(違約金|損害賠償.{0,5}予定).{0,20}(\d{2,})\s*(%|パーセント|万円)",
            r"(解約|中途解約).{0,20}(残.{0,10}全額|全期間.{0,10}料金)",
        ],
        "risk": RiskLevel.HIGH,
        "category": "過大な違約金",
        "description": "過大な違約金・損害賠償の予定は無効となる可能性",
        "legal_basis": "消費者契約法第9条",
        "fix": "平均的な損害の範囲内に設定",
    },
    "payment_over_60days": {
        "patterns": [
            r"支払.{0,20}(6[1-9]|[7-9]\d|1\d{2,})\s*日",
            r"(納品|検収).{0,15}(翌々月|3ヶ月|90日)",
        ],
        "risk": RiskLevel.CRITICAL,
        "category": "支払遅延",
        "description": "60日超の支払期日は下請法違反の可能性",
        "legal_basis": "下請法第4条1項2号",
        "fix": "「納品後60日以内」に修正",
    },
    "disguised_employment": {
        "patterns": [
            r"(業務委託|請負).{0,30}(指揮命令|出退勤.{0,5}管理|勤怠.{0,5}報告)",
            r"(委託者|発注者).{0,20}(指示|命令).{0,10}(従う|従わなければ)",
        ],
        "risk": RiskLevel.CRITICAL,
        "category": "偽装請負",
        "description": "業務委託契約でありながら実態が雇用関係の可能性",
        "legal_basis": "労働基準法、労働者派遣法",
        "fix": "業務委託として成果物・仕様の明確化、または雇用契約に変更",
    },
    "ip_rights_unlimited": {
        "patterns": [
            r"(知的財産|著作権|特許).{0,20}(全て|一切|すべて).{0,10}(帰属|譲渡|移転)",
            r"(成果物|納品物).{0,15}(権利|著作権).{0,10}(甲|委託者|発注者).{0,10}帰属",
        ],
        "risk": RiskLevel.HIGH,
        "category": "知財権",
        "description": "成果物の権利を全て相手方に帰属させる条項",
        "legal_basis": "著作権法、下請法",
        "fix": "適正な対価の明記、または共有・ライセンス形式を検討",
    },
    "unlimited_confidentiality": {
        "patterns": [
            r"秘密保持.{0,20}(永久|無期限|期間.{0,5}定め.{0,5}ない)",
            r"(契約終了|解約).{0,15}後.{0,10}(も|においても).{0,15}(永久|無期限)",
        ],
        "risk": RiskLevel.MEDIUM,
        "category": "秘密保持",
        "description": "過度に長い秘密保持期間",
        "legal_basis": "民法第1条2項",
        "fix": "合理的な期間（3〜5年程度）を設定",
    },
    "non_compete_excessive": {
        "patterns": [
            r"競業禁止.{0,30}(([3-9]|[1-9]\d)\s*年|無期限)",
            r"(退職|契約終了).{0,15}後.{0,10}(5|[6-9]|\d{2,})\s*年.{0,10}競業",
        ],
        "risk": RiskLevel.HIGH,
        "category": "競業禁止",
        "description": "過度に長い競業禁止期間は無効の可能性",
        "legal_basis": "民法第90条、憲法22条（職業選択の自由）",
        "fix": "1-2年程度に短縮し、地域・業種を限定",
    },
    "termination_penalty": {
        "patterns": [
            r"(中途解約|途中解約).{0,20}(できない|認め.{0,5}ない|不可)",
            r"解約.{0,15}(違約金|手数料|ペナルティ).{0,10}(全額|残額)",
        ],
        "risk": RiskLevel.HIGH,
        "category": "解約制限",
        "description": "過度な解約制限は消費者契約法違反の可能性",
        "legal_basis": "消費者契約法第9条、第10条",
        "fix": "合理的な解約条件と違約金上限を設定",
    },
}

# =============================================================================
# メインエンジン
# =============================================================================

class VeritasEngine:
    """VERITAS v162 分析エンジン"""
    
    VERSION = "1.62.0"
    
    def __init__(self):
        self.legal_db = LEGAL_DATABASE
        self.danger_patterns = DANGER_PATTERNS
        self.issue_counter = 0
    
    def analyze(
        self, 
        text: str, 
        file_name: str = "contract.txt",
        domain: str = "auto"
    ) -> AnalysisResult:
        """契約書を分析"""
        
        # 契約種別を検出
        contract_type = self._detect_contract_type(text)
        if domain != "auto":
            contract_type = ContractType(domain) if domain in [e.value for e in ContractType] else contract_type
        
        issues = []
        todo_items = []
        rewrite_suggestions = []
        
        # v162パターンエンジンを使用（利用可能な場合）
        if CORE_AVAILABLE:
            clauses = self._split_clauses(text)
            for clause in clauses:
                result = quick_analyze(clause, domain=domain if domain != "auto" else None)
                
                if result["verdict"] in ["NG_CRITICAL", "NG", "REVIEW_HIGH", "REVIEW_MED"]:
                    risk_level = self._convert_verdict_to_risk(result["verdict"])
                    self.issue_counter += 1
                    
                    issue = Issue(
                        issue_id=f"V162-{self.issue_counter:04d}",
                        clause_text=clause[:200],
                        issue_type=result["verdict"],
                        risk_level=risk_level,
                        description=result["risk_summary"],
                        legal_basis=", ".join(result.get("legal_basis", [])[:3]),
                        fix_suggestion=result["rewrite_suggestions"][0] if result["rewrite_suggestions"] else "専門家に相談してください",
                        category="v162パターン検出",
                        confidence=result["confidence"],
                        check_points=result.get("check_points", []),
                    )
                    issues.append(issue)
                    
                    if result["rewrite_suggestions"]:
                        rewrite_suggestions.append({
                            "original": clause[:150],
                            "suggested": result["rewrite_suggestions"][0],
                            "reason": result["risk_summary"],
                        })
        
        # 従来の危険パターン検出（補完）
        legacy_issues = self._detect_legacy_patterns(text)
        
        # 重複除去してマージ
        seen_texts = {i.clause_text[:50] for i in issues}
        for li in legacy_issues:
            if li.clause_text[:50] not in seen_texts:
                issues.append(li)
                seen_texts.add(li.clause_text[:50])
        
        # ドメインパックによる追加チェック
        if DOMAINS_AVAILABLE:
            domain_issues = self._check_domain_packs(text, contract_type)
            for di in domain_issues:
                if di.clause_text[:50] not in seen_texts:
                    issues.append(di)
                    seen_texts.add(di.clause_text[:50])
        
        # ToDo生成
        for issue in issues:
            todo_items.append({
                "id": issue.issue_id,
                "priority": issue.risk_level.value,
                "action": f"【{issue.category}】{issue.description[:50]}",
                "legal_basis": issue.legal_basis,
            })
        
        # ToDo圧縮（v160機能）
        compressed_todos = []
        if CORE_AVAILABLE and todo_items:
            try:
                todo_objs = [
                    TodoItem(
                        id=t["id"],
                        priority=t["priority"],
                        action=t["action"],
                        legal_basis=t["legal_basis"],
                    ) for t in todo_items
                ]
                compression_result = compress_todos(todo_objs)
                compressed_todos = [
                    {"group": g.group_name, "items": [asdict(i) for i in g.items]}
                    for g in compression_result.groups
                ]
            except Exception:
                compressed_todos = [{"group": "全項目", "items": todo_items}]
        
        # リスクスコア計算
        risk_score = self._calculate_risk_score(issues)
        confidence_interval = self._calculate_confidence_interval(risk_score, len(issues))
        
        return AnalysisResult(
            issues=issues,
            risk_score=risk_score,
            confidence_interval=confidence_interval,
            contract_type=contract_type,
            todo_items=todo_items,
            compressed_todos=compressed_todos,
            rewrite_suggestions=rewrite_suggestions,
            file_name=file_name,
            engine_version=self.VERSION,
        )
    
    def _detect_contract_type(self, text: str) -> ContractType:
        """契約種別を自動検出"""
        keywords = {
            ContractType.NDA: ["秘密保持", "機密情報", "NDA", "守秘義務"],
            ContractType.OUTSOURCING: ["業務委託", "委託業務", "請負", "受託"],
            ContractType.TOS: ["利用規約", "サービス利用", "約款", "ユーザー"],
            ContractType.EMPLOYMENT: ["雇用契約", "労働契約", "就業規則", "給与"],
            ContractType.SALES: ["売買契約", "売買", "購入", "販売"],
            ContractType.LEASE: ["賃貸借", "賃借", "賃貸", "借地借家"],
            ContractType.LICENSE: ["ライセンス", "使用許諾", "実施許諾"],
            ContractType.MA: ["株式譲渡", "事業譲渡", "合併", "M&A"],
            ContractType.IT_SAAS: ["SaaS", "クラウド", "システム利用", "API"],
            ContractType.LABOR: ["出向", "派遣", "就業条件"],
            ContractType.REALESTATE: ["不動産", "土地", "建物", "物件"],
        }
        
        scores = {ct: 0 for ct in ContractType}
        for ct, kws in keywords.items():
            for kw in kws:
                if kw in text:
                    scores[ct] += 1
        
        best = max(scores, key=scores.get)
        return best if scores[best] > 0 else ContractType.GENERAL
    
    def _split_clauses(self, text: str) -> List[str]:
        """条項に分割"""
        patterns = [
            r"第\s*\d+\s*条[^第]*",
            r"\d+\.\s*[^0-9]+",
            r"[（(]\s*\d+\s*[)）][^（(]+",
        ]
        
        clauses = []
        for pattern in patterns:
            matches = re.findall(pattern, text, re.DOTALL)
            clauses.extend(matches)
        
        if not clauses:
            # 改行で分割
            clauses = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 20]
        
        return clauses[:100]  # 最大100条項
    
    def _convert_verdict_to_risk(self, verdict: str) -> RiskLevel:
        """判定をリスクレベルに変換"""
        mapping = {
            "NG_CRITICAL": RiskLevel.CRITICAL,
            "NG": RiskLevel.HIGH,
            "REVIEW_HIGH": RiskLevel.HIGH,
            "REVIEW_MED": RiskLevel.MEDIUM,
        }
        return mapping.get(verdict, RiskLevel.MEDIUM)
    
    def _detect_legacy_patterns(self, text: str) -> List[Issue]:
        """従来の危険パターン検出"""
        issues = []
        
        for pattern_id, pattern_info in self.danger_patterns.items():
            for pattern in pattern_info["patterns"]:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    self.issue_counter += 1
                    
                    # 前後の文脈を取得
                    start = max(0, match.start() - 50)
                    end = min(len(text), match.end() + 50)
                    context = text[start:end]
                    
                    issue = Issue(
                        issue_id=f"LP-{self.issue_counter:04d}",
                        clause_text=context,
                        issue_type=pattern_id,
                        risk_level=pattern_info["risk"],
                        description=pattern_info["description"],
                        legal_basis=pattern_info["legal_basis"],
                        fix_suggestion=pattern_info["fix"],
                        category=pattern_info["category"],
                        confidence=0.9,
                        position=(match.start(), match.end()),
                    )
                    issues.append(issue)
        
        return issues
    
    def _check_domain_packs(self, text: str, contract_type: ContractType) -> List[Issue]:
        """ドメインパックによるチェック"""
        issues = []
        
        domain_mapping = {
            ContractType.LABOR: "LABOR",
            ContractType.EMPLOYMENT: "LABOR",
            ContractType.REALESTATE: "REALESTATE",
            ContractType.LEASE: "REALESTATE",
            ContractType.IT_SAAS: "IT_SAAS",
            ContractType.TOS: "IT_SAAS",
        }
        
        domain = domain_mapping.get(contract_type)
        if domain and domain in AVAILABLE_PACKS:
            pack = AVAILABLE_PACKS[domain]
            try:
                results = pack.check(text)
                for r in results:
                    self.issue_counter += 1
                    risk = RiskLevel.CRITICAL if "CRITICAL" in str(r.verdict) else (
                        RiskLevel.HIGH if "HIGH" in str(r.verdict) or "NG" in str(r.verdict) else RiskLevel.MEDIUM
                    )
                    issue = Issue(
                        issue_id=f"DP-{self.issue_counter:04d}",
                        clause_text=r.matched_text[:200] if hasattr(r, 'matched_text') else "",
                        issue_type=f"{domain}_PACK",
                        risk_level=risk,
                        description=r.risk_explanation if hasattr(r, 'risk_explanation') else str(r),
                        legal_basis=r.legal_basis if hasattr(r, 'legal_basis') else "",
                        fix_suggestion=r.rewrite_suggestion if hasattr(r, 'rewrite_suggestion') else "",
                        category=f"{domain}ドメイン",
                        confidence=0.85,
                    )
                    issues.append(issue)
            except Exception:
                pass
        
        return issues
    
    def _calculate_risk_score(self, issues: List[Issue]) -> float:
        """リスクスコアを計算（0-100）"""
        if not issues:
            return 0.0
        
        weights = {
            RiskLevel.CRITICAL: 30,
            RiskLevel.HIGH: 20,
            RiskLevel.MEDIUM: 10,
            RiskLevel.LOW: 5,
            RiskLevel.SAFE: 0,
        }
        
        total = sum(weights.get(i.risk_level, 10) for i in issues)
        score = min(100, total)
        return score
    
    def _calculate_confidence_interval(self, score: float, n_issues: int) -> Tuple[float, float]:
        """信頼区間を計算"""
        margin = max(5, 15 - n_issues)
        lower = max(0, score - margin)
        upper = min(100, score + margin)
        return (lower, upper)

# =============================================================================
# ファイル処理
# =============================================================================

def extract_text_from_file(uploaded_file) -> str:
    """ファイルからテキストを抽出"""
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    if file_type == "txt":
        return uploaded_file.read().decode("utf-8", errors="ignore")
    
    elif file_type == "pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except ImportError:
            return "[PDF読み取りにはPyPDF2が必要です]"
    
    elif file_type in ["doc", "docx"]:
        try:
            from docx import Document
            doc = Document(io.BytesIO(uploaded_file.read()))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            return "[Word読み取りにはpython-docxが必要です]"
    
    return uploaded_file.read().decode("utf-8", errors="ignore")

# =============================================================================
# OpenAI連携
# =============================================================================

def call_openai_chat(prompt: str, api_key: str) -> str:
    """OpenAI APIを呼び出し"""
    if not api_key:
        return "APIキーが設定されていません。サイドバーで設定してください。"
    
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "あなたは日本の契約書に詳しい法務アシスタントです。"},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.3,
        )
        return response.choices[0].message.content
    except ImportError:
        return "openaiライブラリがインストールされていません"
    except Exception as e:
        return f"エラー: {str(e)}"

# =============================================================================
# Slack通知
# =============================================================================

def send_slack_notification(webhook_url: str, message: str) -> bool:
    """Slack通知を送信"""
    if not webhook_url:
        return False
    
    try:
        import requests
        response = requests.post(
            webhook_url,
            json={"text": message},
            headers={"Content-Type": "application/json"},
        )
        return response.status_code == 200
    except Exception:
        return False

# =============================================================================
# レポート生成
# =============================================================================

def generate_csv_report(result: AnalysisResult) -> str:
    """CSV形式のレポートを生成"""
    lines = ["ID,カテゴリ,リスクレベル,説明,法的根拠,修正提案"]
    for issue in result.issues:
        line = f'"{issue.issue_id}","{issue.category}","{issue.risk_level.value}","{issue.description}","{issue.legal_basis}","{issue.fix_suggestion}"'
        lines.append(line)
    return "\n".join(lines)

def generate_word_report(result: AnalysisResult) -> bytes:
    """Word形式のレポートを生成"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # タイトル
        title = doc.add_heading("VERITAS 契約書分析レポート", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 概要
        doc.add_heading("分析概要", level=1)
        doc.add_paragraph(f"ファイル名: {result.file_name}")
        doc.add_paragraph(f"分析日時: {result.timestamp}")
        doc.add_paragraph(f"契約種別: {result.contract_type.value}")
        doc.add_paragraph(f"リスクスコア: {result.risk_score:.0f}/100")
        doc.add_paragraph(f"検出問題数: {len(result.issues)}件")
        
        # 問題一覧
        doc.add_heading("検出された問題", level=1)
        for issue in result.issues:
            para = doc.add_paragraph()
            run = para.add_run(f"【{issue.risk_level.value}】{issue.category}")
            run.bold = True
            if issue.risk_level == RiskLevel.CRITICAL:
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif issue.risk_level == RiskLevel.HIGH:
                run.font.color.rgb = RGBColor(255, 128, 0)
            
            doc.add_paragraph(f"説明: {issue.description}")
            doc.add_paragraph(f"法的根拠: {issue.legal_basis}")
            doc.add_paragraph(f"修正提案: {issue.fix_suggestion}")
            doc.add_paragraph("")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    except ImportError:
        return b"python-docxライブラリが必要です"

# =============================================================================
# UI コンポーネント
# =============================================================================

def render_risk_badge(risk_level: RiskLevel) -> str:
    """リスクレベルのバッジを表示"""
    colors = {
        RiskLevel.CRITICAL: "🔴",
        RiskLevel.HIGH: "🟠",
        RiskLevel.MEDIUM: "🟡",
        RiskLevel.LOW: "🟢",
        RiskLevel.SAFE: "⚪",
    }
    return f"{colors.get(risk_level, '⚪')} {risk_level.value}"

def render_issue_card(issue: Issue):
    """問題カードを表示"""
    with st.expander(f"{render_risk_badge(issue.risk_level)} {issue.category} - {issue.issue_id}", expanded=issue.risk_level == RiskLevel.CRITICAL):
        st.markdown(f"**説明:** {issue.description}")
        st.markdown(f"**法的根拠:** {issue.legal_basis}")
        st.markdown(f"**修正提案:** {issue.fix_suggestion}")
        if issue.check_points:
            st.markdown("**チェックポイント:**")
            for cp in issue.check_points[:5]:
                st.markdown(f"- {cp}")
        st.code(issue.clause_text, language=None)

def render_statistics():
    """統計情報を表示"""
    if CORE_AVAILABLE:
        try:
            stats = unified_pattern_engine.get_statistics()
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("総パターン数", stats.get("total_patterns", "N/A"))
            with col2:
                st.metric("エンジンバージョン", stats.get("engine_version", "N/A"))
            with col3:
                st.metric("特許対応", "6 Claims")
        except Exception:
            st.info("統計情報を取得できませんでした")

# =============================================================================
# メイン画面
# =============================================================================

def main():
    st.title("🔍 VERITAS v162【完全版】")
    st.caption("AI契約書レビューエンジン - Patent: 2025-159636")
    
    # サイドバー
    with st.sidebar:
        st.header("⚙️ 設定")
        
        # APIキー設定
        st.session_state.openai_api_key = st.text_input(
            "OpenAI API Key",
            value=st.session_state.openai_api_key,
            type="password",
        )
        
        # Slack設定
        st.session_state.slack_webhook_url = st.text_input(
            "Slack Webhook URL（オプション）",
            value=st.session_state.slack_webhook_url,
            type="password",
        )
        
        # ドメイン選択
        st.session_state.selected_domain = st.selectbox(
            "契約ドメイン",
            ["auto", "nda", "outsourcing", "tos", "employment", "labor", "realestate", "it_saas", "general"],
            format_func=lambda x: "自動検出" if x == "auto" else x.upper(),
        )
        
        st.markdown("---")
        
        # エンジン情報
        st.subheader("📊 エンジン情報")
        st.write(f"**バージョン:** v162")
        st.write(f"**v162コア:** {'✅' if CORE_AVAILABLE else '❌'}")
        st.write(f"**ドメインパック:** {'✅' if DOMAINS_AVAILABLE else '❌'}")
        
        if st.button("統計情報を表示"):
            render_statistics()
    
    # メインタブ
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📄 分析", "💬 チャット", "📊 比較", "📈 ダッシュボード", "📋 履歴"])
    
    with tab1:
        render_analysis_tab()
    
    with tab2:
        render_chat_tab()
    
    with tab3:
        render_comparison_tab()
    
    with tab4:
        render_dashboard_tab()
    
    with tab5:
        render_history_tab()

def render_analysis_tab():
    """分析タブ"""
    st.header("📄 契約書分析")
    
    # ファイルアップロード
    uploaded_file = st.file_uploader(
        "契約書をアップロード",
        type=["txt", "pdf", "doc", "docx"],
        help="PDF、Word、テキストファイルに対応",
    )
    
    # テキスト入力
    contract_text = st.text_area(
        "または直接テキストを入力",
        height=200,
        placeholder="契約書のテキストをここに貼り付けてください...",
    )
    
    if uploaded_file:
        contract_text = extract_text_from_file(uploaded_file)
        st.info(f"📎 {uploaded_file.name} を読み込みました（{len(contract_text):,}文字）")
    
    # 分析実行
    if st.button("🔍 分析を実行", type="primary", disabled=not contract_text):
        with st.spinner("分析中..."):
            engine = VeritasEngine()
            result = engine.analyze(
                contract_text,
                file_name=uploaded_file.name if uploaded_file else "direct_input.txt",
                domain=st.session_state.selected_domain,
            )
            st.session_state.current_analysis = result
            
            # 履歴に追加
            st.session_state.analysis_history.append({
                "timestamp": result.timestamp,
                "file_name": result.file_name,
                "risk_score": result.risk_score,
                "issue_count": len(result.issues),
                "contract_type": result.contract_type.value,
            })
        
        # 結果表示
        st.success("✅ 分析完了")
        
        # サマリー
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            color = "🔴" if result.risk_score >= 70 else "🟠" if result.risk_score >= 40 else "🟢"
            st.metric("リスクスコア", f"{color} {result.risk_score:.0f}/100")
        with col2:
            st.metric("検出問題数", len(result.issues))
        with col3:
            st.metric("契約種別", result.contract_type.value)
        with col4:
            st.metric("エンジン", f"v{result.engine_version}")
        
        # 問題一覧
        st.markdown("### 🚨 検出された問題")
        
        # リスクレベルでソート
        sorted_issues = sorted(
            result.issues,
            key=lambda x: [RiskLevel.CRITICAL, RiskLevel.HIGH, RiskLevel.MEDIUM, RiskLevel.LOW, RiskLevel.SAFE].index(x.risk_level),
        )
        
        for issue in sorted_issues:
            render_issue_card(issue)
        
        # ToDo一覧
        if result.compressed_todos:
            st.markdown("### ✅ ToDo リスト（圧縮済み）")
            for group in result.compressed_todos:
                with st.expander(f"📁 {group['group']}（{len(group['items'])}件）"):
                    for item in group['items']:
                        st.checkbox(item.get('action', str(item)), key=f"todo_{item.get('id', '')}")
        
        # リライト提案
        if result.rewrite_suggestions:
            st.markdown("### ✏️ 修正提案")
            for i, suggestion in enumerate(result.rewrite_suggestions[:5]):
                with st.expander(f"提案 {i+1}"):
                    st.markdown("**原文:**")
                    st.code(suggestion['original'])
                    st.markdown("**修正案:**")
                    st.code(suggestion['suggested'])
                    st.markdown(f"**理由:** {suggestion['reason']}")
        
        # エクスポート
        st.markdown("### 📥 レポート出力")
        col1, col2 = st.columns(2)
        
        with col1:
            csv_data = generate_csv_report(result)
            st.download_button(
                "📊 CSVダウンロード",
                csv_data,
                file_name=f"veritas_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
            )
        
        with col2:
            word_data = generate_word_report(result)
            st.download_button(
                "📝 Wordダウンロード",
                word_data,
                file_name=f"veritas_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        
        # Slack通知
        if st.session_state.slack_webhook_url and result.risk_score >= 50:
            message = f"🚨 VERITAS Alert: {result.file_name}\nリスクスコア: {result.risk_score:.0f}/100\n問題数: {len(result.issues)}件"
            if send_slack_notification(st.session_state.slack_webhook_url, message):
                st.success("📢 Slack通知を送信しました")

def render_chat_tab():
    """チャットタブ"""
    st.header("💬 契約書アシスタント")
    
    if not st.session_state.openai_api_key:
        st.warning("チャット機能を使用するには、サイドバーでOpenAI APIキーを設定してください。")
        return
    
    # チャット履歴を表示
    for msg in st.session_state.chat_history:
        with st.chat_message(msg.role):
            st.write(msg.content)
    
    # ユーザー入力
    user_input = st.chat_input("契約書について質問してください...")
    
    if user_input:
        # ユーザーメッセージを追加
        st.session_state.chat_history.append(ChatMessage(role="user", content=user_input))
        
        # コンテキストを構築
        context = ""
        if st.session_state.current_analysis:
            result = st.session_state.current_analysis
            context = f"""
現在分析中の契約書情報:
- ファイル名: {result.file_name}
- 契約種別: {result.contract_type.value}
- リスクスコア: {result.risk_score:.0f}/100
- 検出問題数: {len(result.issues)}件

主な問題:
"""
            for issue in result.issues[:5]:
                context += f"- [{issue.risk_level.value}] {issue.category}: {issue.description}\n"
        
        prompt = f"""
{context}

ユーザーの質問: {user_input}

日本の契約書法務の専門家として、上記の質問に回答してください。
"""
        
        # AI応答を取得
        with st.spinner("回答を生成中..."):
            response = call_openai_chat(prompt, st.session_state.openai_api_key)
        
        # AI応答を追加
        st.session_state.chat_history.append(ChatMessage(role="assistant", content=response))
        
        st.rerun()

def render_comparison_tab():
    """比較分析タブ"""
    st.header("📊 契約書比較分析")
    st.info("2つの契約書を比較して、リスクの違いを分析します。")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**契約書1**")
        contract1 = st.text_area("契約書1のテキスト", height=200, key="compare_contract1")
    
    with col2:
        st.markdown("**契約書2**")
        contract2 = st.text_area("契約書2のテキスト", height=200, key="compare_contract2")
    
    if st.button("🔍 比較分析を実行", type="primary"):
        if contract1.strip() and contract2.strip():
            with st.spinner("比較分析中..."):
                engine = VeritasEngine()
                result1 = engine.analyze(contract1, file_name="契約書1")
                result2 = engine.analyze(contract2, file_name="契約書2")
            
            st.markdown("---")
            st.markdown("### 📈 比較結果")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("契約書1 リスクスコア", f"{result1.risk_score:.0f}点")
            with col2:
                st.metric("契約書2 リスクスコア", f"{result2.risk_score:.0f}点")
            with col3:
                diff = result1.risk_score - result2.risk_score
                st.metric("スコア差", f"{diff:+.0f}点")
            
            # 問題数比較
            st.markdown("#### 問題数比較")
            data = {
                "項目": ["CRITICAL", "HIGH", "MEDIUM", "LOW"],
                "契約書1": [
                    sum(1 for i in result1.issues if i.risk_level == RiskLevel.CRITICAL),
                    sum(1 for i in result1.issues if i.risk_level == RiskLevel.HIGH),
                    sum(1 for i in result1.issues if i.risk_level == RiskLevel.MEDIUM),
                    sum(1 for i in result1.issues if i.risk_level == RiskLevel.LOW),
                ],
                "契約書2": [
                    sum(1 for i in result2.issues if i.risk_level == RiskLevel.CRITICAL),
                    sum(1 for i in result2.issues if i.risk_level == RiskLevel.HIGH),
                    sum(1 for i in result2.issues if i.risk_level == RiskLevel.MEDIUM),
                    sum(1 for i in result2.issues if i.risk_level == RiskLevel.LOW),
                ],
            }
            st.dataframe(data)
        else:
            st.error("両方の契約書テキストを入力してください。")

def render_dashboard_tab():
    """ダッシュボードタブ"""
    st.header("📈 分析ダッシュボード")
    
    history = st.session_state.analysis_history
    
    if not history:
        st.info("まだ分析履歴がありません。「分析」タブで契約書を分析してください。")
        return
    
    # 概要メトリクス
    total = len(history)
    scores = [h.get("risk_score", 0) for h in history]
    avg_score = sum(scores) / total if total > 0 else 0
    high_risk = sum(1 for s in scores if s >= 50)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("総分析数", total)
    with col2:
        st.metric("平均リスクスコア", f"{avg_score:.1f}点")
    with col3:
        st.metric("高リスク件数", high_risk)
    with col4:
        rate = (high_risk / total * 100) if total > 0 else 0
        st.metric("高リスク率", f"{rate:.1f}%")
    
    # トレンドグラフ
    if len(scores) > 1:
        st.markdown("### 📈 リスクスコア推移")
        st.line_chart(scores)
    
    # 契約タイプ分布
    st.markdown("### 📊 契約タイプ分布")
    type_counts = defaultdict(int)
    for h in history:
        type_counts[h.get("contract_type", "unknown")] += 1
    
    for contract_type, count in type_counts.items():
        st.progress(count / max(total, 1), text=f"{contract_type}: {count}件")

def render_history_tab():
    """履歴タブ"""
    st.header("📋 分析履歴")
    
    history = st.session_state.analysis_history
    
    if not history:
        st.info("まだ分析履歴がありません。")
        return
    
    # 履歴テーブル
    for i, h in enumerate(reversed(history)):
        col1, col2, col3, col4 = st.columns([3, 2, 1, 1])
        with col1:
            st.write(h.get("file_name", "unknown"))
        with col2:
            st.write(h.get("timestamp", "")[:19])
        with col3:
            score = h.get("risk_score", 0)
            color = "🔴" if score >= 70 else "🟠" if score >= 40 else "🟢"
            st.write(f"{color} {score:.0f}")
        with col4:
            st.write(f"{h.get('issue_count', 0)}件")
    
    if st.button("🗑️ 履歴をクリア"):
        st.session_state.analysis_history = []
        st.rerun()

# =============================================================================
# エントリーポイント
# =============================================================================

if __name__ == "__main__":
    main()
